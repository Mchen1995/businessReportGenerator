from docx.shared import Inches
from docx import Document
import datetime
import xlrd
import os


# 读取需求 Excel
def getDemandSheet(filename_of_demands):
    demand_sheet = xlrd.open_workbook(filename_of_demands).sheet_by_index(0)
    count = demand_sheet.nrows - 1  # 需求个数
    return demand_sheet, count


# 读取测试用例 Excel
def getCaseSheet(filename_of_cases):
    sheets_list = []
    wb = xlrd.open_workbook(filename_of_cases)
    num_of_sheets = len(wb.sheet_names())
    for i in range(num_of_sheets):
        sheet = xlrd.open_workbook(filename_of_cases).sheet_by_index(i)
        sheets_list.append(sheet)
    return sheets_list, wb.sheet_names()


# 读取指定目录下测试人员的测试案例（只读取一个文件）
def searchCase():
    file_path = '/Users/chenmin/PycharmProjects/businessReport/'
    listdir = os.listdir(file_path)
    file_found = []
    report_name = ''
    report_file = ''
    author = ''
    title_idx = ''  # 标题索引
    case_idx = ''  # 用例编号
    for filename in listdir:
        if filename.startswith('cases'):
            file_found.append(filename)
            left_of_title_index = filename.index('[')  # [ 位置
            right_of_title_index = filename.index(']')  # ] 位置

            left_of_case_index = filename.index('{')  # { 位置
            right_of_case_index = filename.index('}')  # } 位置

            author = filename[right_of_case_index + 2: -4]
            report_name = '测试报告_' + author
            report_file = report_name + '.docx'
            title_idx = filename[left_of_title_index + 1:right_of_title_index]
            case_idx = filename[left_of_case_index + 1:right_of_case_index]
            break
    return file_found[0], report_name, report_file, author, int(title_idx), int(case_idx)


# 创建文档
def createDoc(head, is_total_file):
    document = Document()
    if is_total_file:
        document.add_heading(head, 0)  # 标题
        document.add_heading('1. 测试日期：', 1)  # 一级标题：测试日期
        document.add_heading('2. 测试人员：黄乃芳、郑杰、纪雅容、贺东琴', 1)  # 一级标题：测试人员
        document.add_heading('3. 测试结果：符合需求，测试通过', 1)  # 一级标题：测试结果
    return document


# 生成参数
def genArgs():
    head = '集团门户业务测试报告_'  # 全文标题
    # filename = 'D:\\tmp\\report.docx'  # 测试报告文件名
    # filename_of_demands = 'D:\\tmp\\demands.xls'  # 需求 Excel

    filename = '/Users/chenmin/PycharmProjects/businessReport/集团门户业务测试报告.docx'  # 测试报告文件名
    filename_of_demands = '/Users/chenmin/PycharmProjects/businessReport/软件下发需求.xls'  # 需求 Excel
    return head, filename, filename_of_demands


# 生成测试结果汇总表格
def createSummarizeTable(document, demand_sheet, count):
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '需求编号'
    hdr_cells[2].text = '测试需求点'
    hdr_cells[3].text = '测试结果'
    for i in range(count):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)  # 序号
        row_cells[1].text = demand_sheet.cell_value(i + 1, 1)  # 需求编号
        row_cells[2].text = demand_sheet.cell_value(i + 1, 2)  # 测试需求点
        row_cells[3].text = '通过'  # 序号
    table.style = 'Colorful List'
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(3)
    for cell in table.columns[2].cells:
        cell.width = Inches(4)
    for cell in table.columns[3].cells:
        cell.width = Inches(1.5)


# 获取子标题的索引，即中间列（4.1级）非空白行的索引汇总
def count_sub_titles(column_data):
    res = []
    for i in range(len(column_data)):
        if column_data[i] != '':
            res.append(i)
    return res


# 获取最低级标题索引（有空白行则结束）
def get_min_titles_index(column_data, begin_index):
    res = []
    for idx in range(begin_index, len(column_data)):
        if column_data[idx] != '':
            res.append(idx)
        else:
            break
    return res


# 生成测试案例段落
def createCaseParagraph(document, author, case_sheet, index_of_case_paragraph, case_begin):
    title = case_sheet.cell_value(0, 0)  # 4.1 xxx
    document.add_heading('4.' + str(index_of_case_paragraph) + ' ' + title, level=2)

    subtitle_gather = case_sheet.col_values(1, 0)  # 中间那列即时子标题（4.1.1 级别）的内容
    min_title_gather = case_sheet.col_values(2, 0)  # 最低级标题（4.1.1.1 级别）的内容
    test_time_gather = case_sheet.col_values(3, 0)  # D 列，测试日期
    pictures_gather = case_sheet.col_values(4, 0)  # E 列，截图文件名
    subtitle_indexes = count_sub_titles(subtitle_gather)  # 4.1.1 级别标题非空白行的索引
    for i in range(len(subtitle_indexes)):
        subtitle_index = subtitle_indexes[i]

        # 子标题，如 4.1.1 新增页签区配置功能
        document.add_heading('4.' + str(index_of_case_paragraph)
                             + '.' + str(i + 1) + ' ' + subtitle_gather[subtitle_index], level=3)
        min_title_indexes = get_min_titles_index(min_title_gather, subtitle_index)
        for j in range(len(min_title_indexes)):
            min_title_index = min_title_indexes[j]
            case_name = min_title_gather[min_title_index]  # 用例名称
            test_date = test_time_gather[min_title_index]  # 测试日期，若为空，则取当前日期
            if test_date == '':
                test_date = '20' + datetime.date.today().strftime("%y%m%d")
            pictures_file = pictures_gather[min_title_index]  # 截图文件名
            document.add_heading('4.' + str(index_of_case_paragraph)
                                 + '.' + str(i + 1)
                                 + '.' + str(j + 1) + ' ' + min_title_gather[min_title_index], level=4)
            case_begin = createCaseTable(document, author, case_begin, case_name, str(test_date), pictures_file)


# 生成测试案例表格
def createCaseTable(document, author, the_case_id, case_name, test_date, pictures_file):
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '测试人'
    hdr_cells[1].text = author
    hdr_cells[2].text = '编写人'
    hdr_cells[3].text = author

    row_cells = table.add_row().cells
    row_cells[0].text = '测试用例编号'
    row_cells[1].text = str(the_case_id)
    row_cells[2].text = '测试日期'
    row_cells[3].text = test_date[0:4] + '年' + test_date[4:6] + '月' + test_date[6:8] + '日'

    row_cells = table.add_row().cells
    row_cells[0].text = '测试用例名称'
    row_cells[1].text = case_name
    table.cell(2, 1).merge(table.cell(2, 2)).merge(table.cell(2, 3))  # 测试用例名称行合并为 2 列

    row_cells = table.add_row().cells
    row_cells[0].text = '测试目标'
    row_cells[1].text = '功能正常'
    table.cell(3, 1).merge(table.cell(3, 2)).merge(table.cell(3, 3))  # 测试目标行合并为 2 列

    row_cells = table.add_row().cells
    row_cells[0].text = '预期结果：功能正常'
    table.cell(4, 0).merge(table.cell(4, 1)).merge(table.cell(4, 2)).merge(table.cell(4, 3))  # 预期结果行合并为 1 列

    row_cells = table.add_row().cells
    row_cells[0].text = '实际结果：【此处应截屏说明】\n'
    table.cell(5, 0).merge(table.cell(5, 1)).merge(table.cell(5, 2)).merge(table.cell(5, 3))  # 实际结果行合并为 1 列
    paragraph = row_cells[0].paragraphs[0]
    picture_list = pictures_file.split(',')
    for pic in picture_list:
        if pic != '':
            run = paragraph.add_run()
            # run.add_picture(pic, width=2800000, height=2800000)
            run.add_picture(pic)

    row_cells = table.add_row().cells
    row_cells[0].text = '结果及意见'
    row_cells[1].text = '测试结果符合需求，测试通过。'
    table.cell(6, 1).merge(table.cell(6, 2)).merge(table.cell(6, 3))  # 结果及意见行合并为 2 列

    row_cells = table.add_row().cells
    row_cells[0].text = '备注'
    row_cells[1].text = ''
    table.cell(7, 1).merge(table.cell(7, 2)).merge(table.cell(7, 3))  # 备注行合并为 2 列

    return the_case_id + 1
