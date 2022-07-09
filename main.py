from docx import Document
from docx.shared import Cm, Inches

# 参数
version = "2.5.7.11"  # 版本号
head = '集团门户业务测试报告_' + version  # 全文标题
filename = head + '.docx'  # 文件名
# 需求
demands = (
    ["Q032L0-XXX", "修复全部删除"],
    ["Q032L0-YYY", "漏洞修复"],
    ["Q032L0-ZZZ", "EAIP地址修改"],
    ["Q032L0-ZZZ", "EAIP地址修改"],
)
count = len(demands)  # 需求个数

# 创建文档
document = Document()
document.add_heading(filename, 0)  # 标题
document.add_heading('1. 测试日期：', 1)  # 一级标题：测试日期
document.add_heading('2. 测试人员：黄乃芳、郑杰、纪雅容、贺东琴', 1)  # 一级标题：测试人员
document.add_heading('3. 测试结果：符合需求，测试通过', 1)  # 一级标题：测试结果
# 测试结果汇总表
table = document.add_table(rows=1, cols=4, style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '需求编号'
hdr_cells[2].text = '测试需求点'
hdr_cells[3].text = '测试结果'
for i in range(count):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i + 1)  # 序号
    row_cells[1].text = demands[i][0]  # 需求编号
    row_cells[2].text = demands[i][1]  # 测试需求点
    row_cells[3].text = '通过'  # 序号
table.style = 'Colorful List'
for cell in table.columns[0].cells:
    cell.width = Inches(0.5)
for cell in table.columns[1].cells:
    cell.width = Inches(3)
for cell in table.columns[2].cells:
    cell.width = Inches(4)
for cell in table.columns[3].cells:
    cell.width = Inches(1.5)

document.add_heading('4. 测试项目内容：', 1)  # 一级标题：测试项目内容
document.add_heading('4.' + str(1) + ' 多元金融管理后台增加页签区配置', level=2)

document.save(filename)
